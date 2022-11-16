from docx import Document
from openpyxl import load_workbook

wb = load_workbook(filename = 'test5.xlsx')

sheet = wb['List']

document = Document('test2.docx')

def docx_replace(object, name, value):
    for paragraph in object.paragraphs:
        if name in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if name in inline[i].text:
                    inline[i].text = inline[i].text.replace(name, value)
    for table in object.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace(cell, name, value)

number = None
number_count = 1
count_page = 0
price = 0
table = document.tables[0]

for count in range(2, 2034):
    if (number != sheet['C' + str(count)].value) or (number == None):
        table.columns[5].cells[10].text = str(price)
        if number != None:
            document.save('test1/test' + str(count_page) + '.docx')
            document = Document('test2.docx')
        table = document.tables[0]
        number_count = 1
        count_page += 1
        price = 0
    table_row = table.rows[number_count]
    if number_count == 1:
        docx_replace(document, 'company', sheet['B' + str(count)].value)
        docx_replace(document, 'adress', sheet['E' + str(count)].value)
        docx_replace(document, 'number', sheet['C' + str(count)].value)
        docx_replace(document, 'name', sheet['D' + str(count)].value)
        number = sheet['C' + str(count)].value
    table_row.cells[0].text = str(number_count)
    table_row.cells[1].text = sheet['G' + str(count)].value
    table_row.cells[2].text = "Шт"
    table_row.cells[3].text = str(sheet['H' + str(count)].value)
    table_row.cells[4].text = str(sheet['I' + str(count)].value)
    table_row.cells[5].text = str(sheet['J' + str(count)].value)
    number_count += 1
    price += sheet['J' + str(count)].value

#document.save('test3.docx')
